import docx
import sys
import os

# Create a dummy docx
doc = docx.Document()
doc.add_paragraph("This is a test paragraph.")
doc.save("test.docx")

print("Created test.docx")

try:
    doc = docx.Document("test.docx")
    text = "\n".join([p.text for p in doc.paragraphs])
    print(f"Read text: {text}")
    if "This is a test paragraph." in text:
        print("SUCCESS")
    else:
        print("FAILURE: Content mismatch")
except Exception as e:
    print(f"FAILURE: {e}")

os.remove("test.docx")
